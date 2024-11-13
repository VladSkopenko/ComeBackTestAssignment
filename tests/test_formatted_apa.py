import unittest
from io import BytesIO

from docx import Document
from docx.shared import Pt

from src.services.formatted_apa import FormatterApa


class TestFormatterApa(unittest.TestCase):
    """
    Клас для тестування форматування документа в стилі APA.
    """

    def setUp(self):
        """
        Ініціалізація перед кожним тестом. Створюємо новий документ з текстом "Hello, World!".
        """
        doc = Document()
        doc.add_paragraph("Hello, World!")

        self.file_content = BytesIO()
        doc.save(self.file_content)
        self.file_content.seek(0)

        self.formatter = FormatterApa(self.file_content.read())

    def test_formatted_font(self):
        """
        Тестуємо правильність встановлення шрифту та міжстрочного інтервалу.
        """
        self.formatter.formatted_font()
        style = self.formatter.doc.styles['Normal']
        self.assertEqual(style.font.name, 'Times New Roman')
        self.assertEqual(style.font.size, Pt(12))
        self.assertEqual(style.paragraph_format.line_spacing, 2)

    def test_formatted_margins(self):
        """
        Тестуємо правильність встановлення полів документа.
        """
        self.formatter.formatted_margins()
        for section in self.formatter.doc.sections:
            self.assertEqual(section.left_margin, Pt(72))
            self.assertEqual(section.right_margin, Pt(72))
            self.assertEqual(section.top_margin, Pt(72))
            self.assertEqual(section.bottom_margin, Pt(72))

    def test_get_formatted_stream(self):
        """
        Тестуємо отримання відформатованого потоку байтів.
        """
        self.formatter.formatted_font()
        self.formatter.formatted_margins()

        formatted_stream = self.formatter.get_formatted_stream()
        self.assertIsInstance(formatted_stream, BytesIO)

        formatted_stream.seek(0)
        doc = Document(formatted_stream)
        self.assertGreater(len(doc.paragraphs), 0)

        self.assertEqual(doc.paragraphs[0].text, "Hello, World!")


if __name__ == '__main__':
    unittest.main()
