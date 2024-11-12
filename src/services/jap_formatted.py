from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class JapFormatter:
    def __init__(self, title, authors, affiliations, abstract, keywords, main_text, references):
        self.doc = Document()
        self.title = title
        self.authors = authors
        self.affiliations = affiliations
        self.abstract = abstract
        self.keywords = keywords
        self.main_text = main_text
        self.references = references

    def format_title_page(self):

        title_para = self.doc.add_paragraph(self.title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(14)

        self.doc.add_paragraph(self.authors, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(self.affiliations, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("Author Note", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(self.authors, style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("Contact details and author note content go here.",
                               style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    def format_abstract(self):
        self.doc.add_paragraph("Abstract", style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
        abstract_para = self.doc.add_paragraph(self.abstract)
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def format_keywords(self):
        self.doc.add_paragraph("Keywords:", style='Normal').alignment = WD_ALIGN_PARAGRAPH.LEFT
        keywords_para = self.doc.add_paragraph(self.keywords)
        keywords_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        keywords_para.paragraph_format.left_indent = Pt(10)

    def format_main_text(self):
        self.doc.add_paragraph(self.title, style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
        for section in self.main_text:
            self.add_section(section)

    def add_section(self, section):
        title = section.get('title')
        content = section.get('content')
        level = section.get('level', 1)

        if level == 1:
            self.doc.add_paragraph(title, style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif level == 2:
            self.doc.add_paragraph(title, style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            self.doc.add_paragraph(title, style='Heading 3').alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.doc.add_paragraph(content)

    def format_references(self):
        self.doc.add_paragraph("References", style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ref in self.references:
            self.doc.add_paragraph(ref)

    def save_document(self, filename="formatted_jap_document.docx"):
        self.doc.save(filename)


# Example usage:
